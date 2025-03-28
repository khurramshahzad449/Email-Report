import boto3
import json
import argparse
import requests
import time
from datetime import datetime
from config import *
from jinja2 import Template
from docx import Document
import hmac
import hashlib
import base64

class ClaudeAnalyzer:
    def __init__(self):
        self.bedrock = boto3.client(
            'bedrock-runtime',
            aws_access_key_id=AWS_ACCESS_KEY_ID,
            aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
            region_name=AWS_REGION,
            verify=False
        )

    def analyze_transcript(self, transcript, ideal_pitch, coaching_guide):
        prompt = f"""You are an expert sales coach analyzing sales call transcripts.
        You should acts as: 
          1. Evaluator of live or recorded sales calls 
          2. Scorer based on structured pitch performance 
          3. Coach that delivers personalized feedback to the rep 
          4. Trainer that surfaces missed messaging, and improves objection handling 

        Your primary goal is to To ensure sales reps consistently pitch EveryAction as the only AI-powered 
        advocacy and fundraising platform built on a nonprofit CRM, and use differentiated messaging to 
        close midmarket nonprofit deals ($3M–$15M orgs). 

        Use the evaluation categories and weighting specified in this coaching guide: {coaching_guide}. 

        The key pitch elements that you should listen for should be extracted from the ideal pitch: {ideal_pitch}.

        These are also called out in the coaching guide. The coaching guide also provides missed oportunities to
        flag and objections to handle. It also provides a grading template which should be used to generate the
        final report.

        Now, let's analyze this call transcript and provide a detailed analysis of these points:
            1. Evaluate this sales call transcript. Grade the rep across messaging accuracy, 
            competitive differentiation, needs discovery, objection handling, and pitch structure. 
            Use ideal pitch messaging framework.
            2. What key product or market points did the rep fail to mention? Suggest where and how they could 
            have improved this pitch using the EveryAction case study or feature set.
            3. Analyze how the rep handled objections. Did they effectively counter concerns about existing tools or 
            integrations? Suggest better phrasing aligned with EveryAction's positioning.

        The call transcript is: {transcript}
        Try to use the EBI framework to provide feedback.  Be positive and constructive. 
        IMPORTANT: Your response must be a valid JSON object with the following structure:
        {{
            "didWell": ["list of strengths with brief explanations"],
            "improvements": ["list of areas for improvement with specific examples"],
            "finalScore": <number between 1 and 10>,
            "coachingTips": ["list of coaching tips"]
        }}
        Do not include any markdown formatting or additional text outside the JSON structure."""

        try:
            response = self.bedrock.invoke_model(
                modelId='anthropic.claude-3-sonnet-20240229-v1:0',
                body=json.dumps({
                    "anthropic_version": "bedrock-2023-05-31",
                    "max_tokens": 3000,
                    "temperature": 0.7,
                    "messages": [{"role": "user", "content": prompt}]
                })
            )
            
            # Parse the response from AWS Bedrock
            response_body = json.loads(response['body'].read())
            
            # Extract the text content from Claude's response
            content = response_body['content'][0]['text']
            
            # Clean up the response text
            # Remove markdown code block markers and any leading/trailing whitespace
            content = content.replace('```json', '').replace('```', '').strip()
            
            # Try to find JSON content within the response
            # Look for content between first { and last }
            start_idx = content.find('{')
            end_idx = content.rfind('}') + 1
            
            if start_idx == -1 or end_idx == 0:
                raise ValueError("No JSON object found in response")
                
            json_str = content[start_idx:end_idx]
            
            # Parse the JSON
            result = json.loads(json_str)
            
            # Validate the required fields
            required_fields = ['didWell', 'improvements', 'finalScore', 'coachingTips']
            for field in required_fields:
                if field not in result:
                    raise ValueError(f"Missing required field: {field}")
                if field == 'finalScore' and not isinstance(result[field], (int, float)):
                    raise ValueError("finalScore must be a number")
                if field in ['didWell', 'improvements', 'coachingTips'] and not isinstance(result[field], list):
                    raise ValueError(f"{field} must be a list")
            
            return result
            
        except json.JSONDecodeError as e:
            print(f"Error parsing JSON response: {e}")
            print(f"Raw response content: {content}")
            raise
        except Exception as e:
            print(f"Error in analyze_transcript: {e}")
            raise

class ReportGenerator:
    def __init__(self):
        self.email_template = Template(EMAIL_TEMPLATE)

    def generate_email(self, analysis_results, call_details):
        # Create template variables
        template_vars = {
            'didWell': '\n'.join(f"- {s}" for s in analysis_results['didWell']),
            'improvements': '\n'.join(f"- {s}" for s in analysis_results['improvements']),
            'finalScore': analysis_results['finalScore'],
            'coachingTips': '\n'.join(f"- {s}" for s in analysis_results['coachingTips'])
        }

        # Render the template
        return self.email_template.render(**template_vars)

def read_transcript(file_path):
    """Read transcript from a file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read().strip()
    except Exception as e:
        print(f"Error reading transcript file: {e}")
        return None

def read_word_document(file_path):
    """Read content from a Word document, including both paragraphs and tables."""
    try:
        doc = Document(file_path)
        full_text = []
        
        # Read paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Read tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:  # Only add non-empty rows
                    full_text.append(" | ".join(row_text))
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error reading Word document: {e}")
        return None

def get_gong_signature(access_key, secret_key, timestamp, request_id, method, path, body=""):
    """Generate signature for Gong API request."""
    string_to_sign = f"{timestamp}\n{request_id}\n{method}\n{path}\n{body}"
    signature = hmac.new(
        secret_key.encode('utf-8'),
        string_to_sign.encode('utf-8'),
        hashlib.sha256
    ).digest()
    return base64.b64encode(signature).decode('utf-8')

def get_gong_user_info(user_id):
    """Fetch user information from Gong API."""
    try:
        timestamp = str(int(time.time() * 1000))
        request_id = str(int(time.time() * 1000))
        path = f"/v2/users/{user_id}"
        method = "GET"
        
        # Generate signature
        signature = get_gong_signature(
            GONG_ACCESS_KEY,
            GONG_SECRET_KEY,
            timestamp,
            request_id,
            method,
            path
        )
        
        # Make request
        headers = {
            "Authorization": f"Basic {base64.b64encode(f'{GONG_ACCESS_KEY}:{GONG_SECRET_KEY}'.encode()).decode()}",
            "X-Gong-Timestamp": timestamp,
            "X-Gong-Request-Id": request_id,
            "X-Gong-Signature": signature,
            "Content-Type": "application/json"
        }
        
        response = requests.get(
            f"{GONG_BASE_URL}{path}",
            headers=headers,
            verify=False
        )
        
        if response.status_code == 200:
            return response.json()
        else:
            print(f"Error fetching user info: {response.status_code} - {response.text}")
            return None
            
    except Exception as e:
        print(f"Error fetching user info: {e}")
        return None

def get_gong_parties(conversation_id):
    """Fetch parties information from Gong API using the extensive endpoint."""
    try:
        timestamp = str(int(time.time() * 1000))
        request_id = str(int(time.time() * 1000))
        path = f"/v2/calls/extensive"
        method = "POST"
        
        # Request body
        body = {
            "filter": {
                "callIds": [conversation_id]
            },
            "contentSelector" : {
                "exposedFields":{
                    "parties": True
                }
            }
        }
        body_json = json.dumps(body)
        
        # Generate signature
        signature = get_gong_signature(
            GONG_ACCESS_KEY,
            GONG_SECRET_KEY,
            timestamp,
            request_id,
            method,
            path,
            body_json
        )
        
        # Make request
        headers = {
            "Authorization": f"Basic {base64.b64encode(f'{GONG_ACCESS_KEY}:{GONG_SECRET_KEY}'.encode()).decode()}",
            "X-Gong-Timestamp": timestamp,
            "X-Gong-Request-Id": request_id,
            "X-Gong-Signature": signature,
            "Content-Type": "application/json"
        }
        
        response = requests.post(
            f"{GONG_BASE_URL}{path}",
            headers=headers,
            json=body,
            verify=False
        )
        
        if response.status_code == 200:
            data = response.json()
            print("--------------------------------")
            print(data)
            print("--------------------------------")
            if data.get('calls') and len(data['calls']) > 0:
                return data['calls'][0].get('parties', [])
        else:
            print(f"Error fetching parties info: {response.status_code} - {response.text}")
        return []
            
    except Exception as e:
        print(f"Error fetching parties info: {e}")
        return []

def fetch_gong_transcript(conversation_id):
    """Fetch transcript from Gong API."""
    try:
        timestamp = str(int(time.time() * 1000))
        request_id = str(int(time.time() * 1000))
        path = "/v2/calls/transcript"
        method = "POST"
        
        # Request body
        body = {
            "filter": {
                "callIds": [conversation_id]
            }
        }
        body_json = json.dumps(body)
        
        # Generate signature
        signature = get_gong_signature(
            GONG_ACCESS_KEY,
            GONG_SECRET_KEY,
            timestamp,
            request_id,
            method,
            path,
            body_json
        )
        
        # Make request
        headers = {
            "Authorization": f"Basic {base64.b64encode(f'{GONG_ACCESS_KEY}:{GONG_SECRET_KEY}'.encode()).decode()}",
            "X-Gong-Timestamp": timestamp,
            "X-Gong-Request-Id": request_id,
            "X-Gong-Signature": signature,
            "Content-Type": "application/json"
        }
        
        response = requests.post(
            f"{GONG_BASE_URL}{path}",
            headers=headers,
            json=body,
            verify=False
        )
        
        if response.status_code == 200:
            try:
                data = response.json()
                transcript = []
                
                # Get the first call transcript
                if data.get('callTranscripts') and len(data['callTranscripts']) > 0:
                    call_transcript = data['callTranscripts'][0]
                    
                    # Get parties information and create speaker name mapping
                    parties = get_gong_parties(conversation_id)
                    speaker_names = {}
                    for party in parties:
                        speaker_id = party.get('speakerId')
                        if speaker_id:
                            speaker_names[speaker_id] = party.get('name', f"User_{speaker_id}")
                    
                    # Process each transcript section
                    for section in call_transcript.get('transcript', []):
                        topic = section.get('topic', '')
                        speaker_id = section.get('speakerId')
                        
                        # Get speaker name from cached mapping or use ID
                        speaker_name = speaker_names.get(speaker_id, f"User_{speaker_id}")
                        
                        for sentence in section.get('sentences', []):
                            text = sentence.get('text', '')
                            if text:
                                # Format: [Topic] Speaker Name: Text
                                transcript.append(f"[{topic}] {speaker_name}: {text}")
                
                return '\n'.join(transcript)
            except json.JSONDecodeError as e:
                print(f"Error parsing JSON response: {e}")
                print(f"Full response content: {response.text}")
                return None
        else:
            print(f"Error fetching Gong transcript: {response.status_code} - {response.text}")
            return None
            
    except Exception as e:
        print(f"Error fetching Gong transcript: {e}")
        return None

def process_transcript(transcript, ideal_pitch, coaching_guide, sales_rep, customer, duration, output_file):
    """Process a single transcript and generate a report."""
    if not transcript:
        print("Failed to read transcript. Skipping...")
        return False

    # Save transcript to a file
    transcript_file = output_file.replace('.txt', '_transcript.txt')
    try:
        with open(transcript_file, 'w', encoding='utf-8') as f:
            f.write(transcript)
        print(f"Transcript saved to {transcript_file}")
    except Exception as e:
        print(f"Error saving transcript to file: {e}")

    # Initialize analyzer and report generator
    claude_analyzer = ClaudeAnalyzer()
    report_generator = ReportGenerator()

    # Call details
    call_details = {
        'date': datetime.now().strftime('%Y-%m-%d'),
        'sales_rep': sales_rep,
        'customer': customer,
        'duration': duration
    }

    print("Analyzing transcript...")
    # Analyze transcript with Claude
    analysis_results = claude_analyzer.analyze_transcript(
        transcript,
        ideal_pitch,
        coaching_guide
    )

    print("Generating report...")
    # Generate email report
    email_content = report_generator.generate_email(
        analysis_results,
        call_details
    )

    # Save the report
    with open(output_file, 'w') as f:
        f.write(email_content)
    
    print(f"Analysis complete! Report saved to {output_file}")
    return True

def main():
    # Set up argument parser
    parser = argparse.ArgumentParser(description='Analyze sales call transcript using AWS Bedrock Claude')
    parser.add_argument('--transcript', type=str, help='Path to the transcript file (required if not using --gong-ids)')
    parser.add_argument('--ideal-pitch', type=str, required=True, help='Path to the Word document containing the ideal pitch')
    parser.add_argument('--coaching-guide', type=str, required=True, help='Path to the Word document containing the coaching guide')
    parser.add_argument('--sales-rep', type=str, required=True, help='Name of the sales representative')
    parser.add_argument('--customer', type=str, required=True, help='Name of the customer')
    parser.add_argument('--duration', type=str, default='45 minutes', help='Duration of the call')
    parser.add_argument('--output', type=str, default='analysis_report.txt', help='Output file path for the report')
    parser.add_argument('--gong-ids', type=str, help='Comma-separated list of Gong conversation IDs to analyze')
    
    args = parser.parse_args()

    # Validate arguments
    if not args.gong_ids and not args.transcript:
        print("Error: Either --transcript or --gong-ids must be provided")
        return

    # Read ideal pitch from Word document
    ideal_pitch = read_word_document(args.ideal_pitch)
    if not ideal_pitch:
        print("Failed to read ideal pitch document. Exiting...")
        return

    # Read coaching guide from Word document
    coaching_guide = read_word_document(args.coaching_guide)
    if not coaching_guide:
        print("Failed to read coaching guide document. Exiting...")
        return

    if args.gong_ids:
        # Process Gong transcripts
        gong_ids = [id.strip() for id in args.gong_ids.split(',')]
        for i, conversation_id in enumerate(gong_ids):
            print(f"\nProcessing Gong transcript {i+1}/{len(gong_ids)} (ID: {conversation_id})...")
            transcript = fetch_gong_transcript(conversation_id)
            output_file = f"analysis_report_{conversation_id}.txt"
            process_transcript(
                transcript,
                ideal_pitch,
                coaching_guide,
                args.sales_rep,
                args.customer,
                args.duration,
                output_file
            )
    else:
        # Process local transcript
        transcript = read_transcript(args.transcript)
        process_transcript(
            transcript,
            ideal_pitch,
            coaching_guide,
            args.sales_rep,
            args.customer,
            args.duration,
            args.output
        )

if __name__ == "__main__":
    main()